"""
Generate two Word documents from the existing Markdown SOP + README, written
to G:\\My Drive\\HJR-Founder-OS\\01-HJR-Global\\accounting\\ so Harrison can
attach them to Justin's onboarding email.

Run once:
    python generate-justin-docs.py

The script auto-installs python-docx if it's missing. Output filenames follow
the naming-conventions.md format (YYYY-MM-DD_entity_kebab-description.ext).
"""

from __future__ import annotations

import subprocess
import sys
from pathlib import Path

# ---------------------------------------------------------------------------
# Auto-install dependency
# ---------------------------------------------------------------------------
try:
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor, Inches, Cm
except ImportError:
    print("Installing python-docx (one-time)...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor, Inches, Cm

# ---------------------------------------------------------------------------
# Output paths
# ---------------------------------------------------------------------------
DRIVE_ROOT = Path(r"G:\My Drive\HJR-Founder-OS\01-HJR-Global\accounting")
README_OUT = DRIVE_ROOT / "2026-05-21_hjrg_accounting-readme.docx"
SOP_OUT = DRIVE_ROOT / "2026-05-21_hjrg_financial-data-sop.docx"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
BASE_FONT = "Calibri"
MONO_FONT = "Consolas"


def _set_cell_shading(cell, fill_hex: str) -> None:
    """Add background shading to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _set_cell_borders(cell, color_hex: str = "BFBFBF") -> None:
    """Light gray borders on all four sides."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), color_hex)
        borders.append(b)
    tc_pr.append(borders)


def init_doc() -> Document:
    """Create a new document with our default styles."""
    doc = Document()

    # Normal style
    normal = doc.styles["Normal"]
    normal.font.name = BASE_FONT
    normal.font.size = Pt(11)

    # Heading styles
    for name, size in (("Heading 1", 20), ("Heading 2", 14), ("Heading 3", 12)):
        s = doc.styles[name]
        s.font.name = BASE_FONT
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)

    # Page margins (1 inch all around, US Letter implicit from python-docx default)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    return doc


def add_para(doc, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic


def add_mixed_para(doc, segments: list[tuple[str, dict]]) -> None:
    """segments: list of (text, {bold,italic,mono}). Renders inline mixed formatting."""
    p = doc.add_paragraph()
    for text, style in segments:
        run = p.add_run(text)
        run.bold = style.get("bold", False)
        run.italic = style.get("italic", False)
        if style.get("mono"):
            run.font.name = MONO_FONT
            run.font.size = Pt(10)


def add_code_block(doc, text: str) -> None:
    """Render a monospaced block (used for the folder tree)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.15)
    # Light gray background via XML
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    p_pr.append(shd)
    run = p.add_run(text)
    run.font.name = MONO_FONT
    run.font.size = Pt(9)


def add_bullet(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def add_bullet_mixed(doc, segments: list[tuple[str, dict]]) -> None:
    p = doc.add_paragraph(style="List Bullet")
    for text, style in segments:
        run = p.add_run(text)
        run.bold = style.get("bold", False)
        run.italic = style.get("italic", False)
        if style.get("mono"):
            run.font.name = MONO_FONT
            run.font.size = Pt(10)


def add_numbered(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)


def add_numbered_mixed(doc, segments: list[tuple[str, dict]]) -> None:
    p = doc.add_paragraph(style="List Number")
    for text, style in segments:
        run = p.add_run(text)
        run.bold = style.get("bold", False)
        run.italic = style.get("italic", False)
        if style.get("mono"):
            run.font.name = MONO_FONT
            run.font.size = Pt(10)


def add_table(doc, headers: list[str], rows: list[list[str]], col_widths: list[float] | None = None) -> None:
    """Add a clean 1-line-bordered table with a shaded header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = False
    table.allow_autofit = False

    # Headers
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        _set_cell_shading(cell, "D9E2F3")
        _set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Body
    for r_i, row in enumerate(rows, start=1):
        body_cells = table.rows[r_i].cells
        for c_i, val in enumerate(row):
            cell = body_cells[c_i]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(val)
            run.font.size = Pt(10)
            _set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Column widths
    if col_widths:
        for row in table.rows:
            for c_i, w in enumerate(col_widths):
                row.cells[c_i].width = Inches(w)


def add_hr(doc) -> None:
    """Horizontal rule (uses bottom border on an empty paragraph)."""
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "BFBFBF")
    bottom.set(qn("w:space"), "1")
    pbdr.append(bottom)
    p_pr.append(pbdr)


# ---------------------------------------------------------------------------
# README content
# ---------------------------------------------------------------------------
def build_readme() -> None:
    doc = init_doc()

    doc.add_heading("HJR Global - Accounting", level=1)
    add_para(
        doc,
        "This folder is the canonical home for all HJR portfolio financial data. "
        "Justin Moran is the steward; Hayden (Visibility CPA) is a contributor with edit access "
        "on specific shared sheets. Authoritative copies live here, not in Dropbox.",
    )

    doc.add_heading("Where things live (at a glance)", level=2)
    items = [
        ("live-sheets/", "Always-current Google Sheets. Edit in place. Examples: weekly cash flow, OSN monthly metrics, intercompany recon, F3E budget, UFL P&L. Never replace these - they are edited continuously."),
        ("monthly-reports/YYYY-MM/", "Month-end close pack snapshots, append-only. Folder name = month of upload; file name = data period + entity + doctype (e.g., 2026-04_f3e_pl.xlsx)."),
        ("source-system-exports/clover/YYYY-MM/", "OSN Clover daily / weekly CSV exports."),
        ("source-system-exports/qbo/YYYY-MM/", "Ad-hoc QBO exports."),
        ("bank-statements/YYYY-MM/", "Chase / BoA / etc. statements per entity per month."),
        ("quarterly/YYYY-Q{1-4}/", "Quarterly aggregations + 941s."),
        ("annual/YYYY/", "Year-end statements + 1099s."),
        ("tax-filings/YYYY/", "Returns, amendments, IRS correspondence."),
        ("visibility-binder/", "LEGACY. Historical archive of Visibility's organizational scheme. Preserve, don't add to."),
    ]
    for path, desc in items:
        add_bullet_mixed(doc, [(path, {"mono": True, "bold": True}), (f"  -  {desc}", {})])

    doc.add_heading("Naming convention (short version)", level=2)
    add_mixed_para(doc, [
        ("Live sheets:  ", {"bold": True}),
        ("{entity-lower}_{topic-kebab}_LIVE.gsheet", {"mono": True}),
        ("  ->  ", {}),
        ("osn_monthly-metrics_LIVE.gsheet", {"mono": True}),
    ])
    add_mixed_para(doc, [
        ("Snapshots:  ", {"bold": True}),
        ("{data-period}_{entity-lower}_{doctype}.{ext}", {"mono": True}),
        ("  ->  ", {}),
        ("2026-04_f3e_pl.xlsx", {"mono": True}),
    ])
    add_para(
        doc,
        "Doctype shortcuts for financial reports: pl (P&L), bs (balance sheet), cf (cash flow), "
        "ar (A/R aging), ap (A/P aging), tb (trial balance).",
    )

    doc.add_heading("Read the full SOP", level=2)
    add_mixed_para(doc, [
        ("For cadence, migration plan, permissions matrix, and anti-patterns, read ", {}),
        ("2026-05-21_hjrg_financial-data-sop.docx", {"mono": True, "bold": True}),
        (" (attached to your kickoff email) or the canonical Markdown copy at ", {}),
        ("_shared/playbooks/justin-financial-data-sop.md", {"mono": True}),
        (".", {}),
    ])

    doc.add_heading("Who reads from here", level=2)
    add_bullet_mixed(doc, [
        ("Cora", {"bold": True}),
        (" (entity-aware Slack bot): reads ", {}),
        ("live-sheets/", {"mono": True}),
        (", ", {}),
        ("monthly-reports/", {"mono": True}),
        (", ", {}),
        ("source-system-exports/clover/", {"mono": True}),
        (" via its service account. Surfaces source-opaque, freshness-labeled answers in Slack - never exposes filenames or links downstream.", {}),
    ])
    add_bullet_mixed(doc, [
        ("Harrison + Justin", {"bold": True}),
        (": full edit.", {}),
    ])
    add_bullet_mixed(doc, [
        ("Hayden", {"bold": True}),
        (" (Visibility CPA): edit on co-maintained ", {}),
        ("_LIVE", {"mono": True}),
        (" sheets; view elsewhere; no access to Lex PHI.", {}),
    ])
    add_bullet(doc, "Nobody else by default.")

    doc.save(str(README_OUT))
    print(f"Wrote {README_OUT}")


# ---------------------------------------------------------------------------
# SOP content
# ---------------------------------------------------------------------------
FOLDER_TREE = """01-HJR-Global/accounting/
+-- README.md                        -- quick reference, points back to this SOP
|
+-- live-sheets/                     -- ALWAYS-CURRENT Google Sheets. Edit in place. Never replace.
|   +-- hjrg_weekly-cash-flow_LIVE.gsheet
|   +-- hjrg_intercompany-recon_LIVE.gsheet
|   +-- f3e_budget_LIVE.gsheet
|   +-- f3e_inventory_LIVE.gsheet
|   +-- osn_monthly-metrics_LIVE.gsheet
|   +-- osn_item-sales_LIVE.gsheet
|   +-- ufl_pl_LIVE.gsheet
|   +-- lex-lts_cash-flow_LIVE.gsheet
|   +-- lex-lbhs_rita-tracking_LIVE.gsheet      (PHI - tighter ACL, see Permissions)
|   +-- lex_tax-resolution-tracker_LIVE.gsheet
|   +-- hjrp_hampton-cams_LIVE.gsheet
|
+-- monthly-reports/                 -- Month-end close pack snapshots, append-only.
|   +-- YYYY-MM/                     -- Folder name = MONTH OF UPLOAD (not data period).
|       +-- YYYY-MM_entity_doctype.xlsx     -- File name = DATA PERIOD + entity + doctype.
|                                     doctype in {pl, bs, cf, ar, ap, tb}
|
+-- source-system-exports/           -- Raw exports from POS / external accounting systems.
|   +-- clover/                      -- OSN Clover exports.
|   |   +-- YYYY-MM/
|   |       +-- YYYY-MM-DD_osn-{gw|gm|gf|vvp}_{daily-sales|inventory|item-mix}.csv
|   +-- qbo/                         -- Ad-hoc QBO exports Justin pulls manually.
|       +-- YYYY-MM/
|           +-- YYYY-MM-DD_entity_qbo-{report-name}.xlsx
|
+-- bank-statements/                 -- Chase / BoA / etc. statements, by month.
|   +-- YYYY-MM/
|       +-- YYYY-MM_entity_bank-{last4}.pdf
|
+-- quarterly/                       -- Quarterly aggregations / 941s.
|   +-- YYYY-Q{1-4}/
|
+-- annual/                          -- Year-end statements / 1099s.
|   +-- YYYY/
|
+-- tax-filings/                     -- Returns + amendments + IRS correspondence.
|   +-- YYYY/
|
+-- visibility-binder/               -- LEGACY. Preserve as historical archive. Do not touch.
"""


def build_sop() -> None:
    doc = init_doc()

    doc.add_heading("Financial Data SOP - Justin", level=1)
    add_para(
        doc,
        "Canonical operating procedure for HJR's financial data home, owned by Justin Moran. "
        "Last revised 2026-05-21.",
        italic=True,
    )
    add_hr(doc)

    # Why this exists
    doc.add_heading("Why this exists", level=2)
    add_para(
        doc,
        "Until now, the freshest copies of HJR's financial files have lived in Hayden's "
        "(Visibility CPA) Dropbox folders. That's wrong for two reasons:",
    )
    add_numbered_mixed(doc, [
        ("Sovereignty. ", {"bold": True}),
        ("HJR's financials sitting in a third-party firm's cloud is a control risk. If the "
         "Visibility relationship pauses, gets cut, or has a security incident, we lose access to our own data.", {}),
    ])
    add_numbered_mixed(doc, [
        ("Reliability for Cora. ", {"bold": True}),
        ("Cora (the entity-aware Slack bot) cannot reliably read from Hayden's Dropbox - "
         "third-party sharing, inconsistent naming, brittle permissions. To answer \"what's our cash position?\" "
         "or \"OSN April P&L\" in Slack, Cora needs a stable, well-named, HJR-owned data home.", {}),
    ])
    add_mixed_para(doc, [
        ("This SOP makes ", {}),
        (r"G:\My Drive\HJR-Founder-OS\01-HJR-Global\accounting\\", {"mono": True}),
        (" the ", {}),
        ("single canonical financial data home", {"bold": True}),
        (". Justin is the steward. Hayden is a contributor with edit access on specific shared files, "
         "but the canonical copies live here, not in Dropbox.", {}),
    ])

    add_hr(doc)

    # Folder structure
    doc.add_heading("The folder structure", level=2)
    add_code_block(doc, FOLDER_TREE)

    add_mixed_para(doc, [
        ("OSN location codes: ", {"bold": True}),
        ("gw", {"mono": True}),
        (" = Gilbert & Warner;  ", {}),
        ("gm", {"mono": True}),
        (" = Gilbert & McKellips;  ", {}),
        ("gf", {"mono": True}),
        (" = Greenfield & 60;  ", {}),
        ("vvp", {"mono": True}),
        (" = Val Vista & Pecos.", {}),
    ])
    add_mixed_para(doc, [
        ("Entity codes ", {"bold": True}),
        ("follow the canonical list in ", {}),
        ("_shared/playbooks/naming-conventions.md", {"mono": True}),
        (". The one not yet in that list but used here: ", {}),
        ("lex-lts", {"mono": True}),
        (" for Lexington Therapies - to be added at next revision.", {}),
    ])

    add_hr(doc)

    # Naming
    doc.add_heading("Naming conventions", level=2)
    add_para(doc, "Two patterns, one rule each.")

    doc.add_heading("Live sheets (always-current Google Sheets)", level=3)
    add_code_block(doc, "{entity-code-lower}_{topic-kebab}_LIVE.gsheet")
    add_mixed_para(doc, [
        ("Example: ", {}),
        ("osn_monthly-metrics_LIVE.gsheet", {"mono": True}),
    ])
    add_mixed_para(doc, [
        ("The ", {}),
        ("_LIVE", {"mono": True}),
        (" suffix is the contract: this file is edited in place, never replaced, never duplicated. "
         "The Google Sheet ID stays stable forever so Cora's code can pin to it.", {}),
    ])

    doc.add_heading("Snapshot files (monthly close packs, bank statements, exports, etc.)", level=3)
    add_code_block(doc, "{data-period}_{entity-code-lower}_{doctype-or-description}.{ext}")
    add_bullet_mixed(doc, [
        ("data-period", {"mono": True, "bold": True}),
        (": when the data is FROM - ", {}),
        ("2026-04", {"mono": True}),
        (" for April 2026 P&L. Use ", {}),
        ("YYYY-MM", {"mono": True}),
        (" for monthly, ", {}),
        ("YYYY-MM-DD", {"mono": True}),
        (" for daily exports.", {}),
    ])
    add_bullet_mixed(doc, [
        ("entity-code-lower", {"mono": True, "bold": True}),
        (": lowercase, kebab-friendly. ", {}),
        ("hjrg", {"mono": True}),
        (", ", {}),
        ("f3e", {"mono": True}),
        (", ", {}),
        ("osn", {"mono": True}),
        (", ", {}),
        ("osn-gw", {"mono": True}),
        (", ", {}),
        ("lex-lbhs", {"mono": True}),
        (", etc.", {}),
    ])
    add_bullet_mixed(doc, [
        ("doctype", {"mono": True, "bold": True}),
        (" for financial reports: ", {}),
        ("pl", {"mono": True}),
        (" (P&L), ", {}),
        ("bs", {"mono": True}),
        (" (balance sheet), ", {}),
        ("cf", {"mono": True}),
        (" (cash flow), ", {}),
        ("ar", {"mono": True}),
        (" (A/R aging), ", {}),
        ("ap", {"mono": True}),
        (" (A/P aging), ", {}),
        ("tb", {"mono": True}),
        (" (trial balance).", {}),
    ])
    add_bullet_mixed(doc, [
        ("description", {"mono": True, "bold": True}),
        (" (free-form, kebab-case) for everything else.", {}),
    ])

    add_para(doc, "Examples (matching what's already in monthly-reports/2026-05/):")
    for ex in (
        "2026-04_f3e_pl.xlsx",
        "2026-04_osn-gw_bs.xlsx",
        "2026-05-20_osn-vvp_daily-sales.csv",
    ):
        add_bullet_mixed(doc, [(ex, {"mono": True})])

    add_mixed_para(doc, [
        ("Don't: ", {"bold": True}),
        ("spaces, mixed case, em-dashes, special characters, ", {}),
        ("Final", {"mono": True}),
        (", ", {}),
        ("v2-FINAL-real", {"mono": True}),
        (", author initials.", {}),
    ])

    add_hr(doc)

    # Cadence table
    doc.add_heading("Justin's cadence", level=2)
    add_table(
        doc,
        headers=["Cadence", "What", "Where", "When"],
        rows=[
            ["Daily", "OSN Clover CSV export per location (4 stores)", "source-system-exports/clover/YYYY-MM/", "EOD or first thing next morning"],
            ["Weekly (Mon AM)", "Refresh hjrg_weekly-cash-flow_LIVE with prior week actuals", "live-sheets/ (edit in place)", "Monday morning"],
            ["Weekly (Mon AM)", "Sanity check: every _LIVE sheet edited within past 7 days", "live-sheets/", "Monday morning"],
            ["Monthly", "Receive Visibility's close pack -> save P&L/BS/CF/AR/AP per entity", "monthly-reports/YYYY-MM/ (upload month)", "When Hayden ships close pack"],
            ["Monthly", "Update affected _LIVE sheets with new month's actuals", "live-sheets/", "Same day as close pack received"],
            ["Monthly", "Save bank statements for each entity", "bank-statements/YYYY-MM/", "When statements arrive"],
            ["Quarterly", "Quarterly aggregations + 941s", "quarterly/YYYY-Q{n}/", "After Visibility's quarterly review"],
            ["Annual", "Year-end statements + 1099s", "annual/YYYY/", "January / February following year"],
            ["As-needed", "Tax filings, amendments, IRS correspondence", "tax-filings/YYYY/", "When the work happens"],
        ],
        col_widths=[1.1, 2.3, 2.0, 1.5],
    )

    add_mixed_para(doc, [
        ("The one universal rule: ", {"bold": True}),
        ("if Hayden sends you a file via Dropbox, email, or any other channel, it gets pulled into Drive ", {}),
        ("same day", {"bold": True}),
        (" with the right name in the right folder. Files don't sit in Dropbox waiting to be moved later.", {}),
    ])

    add_hr(doc)

    # Migration
    doc.add_heading("One-time migration plan (~2-3 hours of Justin's time)", level=2)
    add_para(doc, "This is the cutover from Dropbox to Drive. Do it in this order:")
    migration_steps = [
        ("Pull every active financial file", " out of Hayden's Dropbox folders into the new Drive paths. Use the naming conventions above as you go."),
        ("For files that are Google Sheets in Dropbox", " (the LIVE sheets Hayden has been maintaining): in Drive, \"Make a copy\" with the new {entity}_{topic}_LIVE.gsheet name into live-sheets/. The new copy gets a new sheet ID - that's fine because Cora reads by the new ID."),
        ("Grant Hayden edit access on the new Drive LIVE sheets", " she'll continue to co-maintain (Weekly Cash Flow, OSN monthly metrics, intercompany-recon, anything else she touches). She loses nothing operationally - she just edits in Drive instead of Dropbox."),
        ("Send Hayden the kickoff note", " (Harrison will forward you a template). Cc Harrison."),
        ("Stop writing to Dropbox", ". Anything new lands in Drive. Existing Dropbox files become read-only historical archives (do not delete - they're audit trail)."),
        ("Confirm Cora's service account has view access", " on all live sheets and the accounting/ folder. The service account email is cora-calendar-sa@<project>.iam.gserviceaccount.com (Harrison has it; he'll send it to you)."),
    ]
    for bold_part, rest in migration_steps:
        add_numbered_mixed(doc, [(bold_part, {"bold": True}), (rest, {})])

    add_hr(doc)

    # Permissions
    doc.add_heading("Permissions matrix", level=2)
    add_table(
        doc,
        headers=["Who", "live-sheets/ (general)", "live-sheets/ (LBHS / Lex PHI)", "monthly-reports/ + snapshots", "visibility-binder/"],
        rows=[
            ["Harrison", "Owner", "Owner", "Owner", "Owner"],
            ["Justin", "Edit", "Edit", "Edit", "Edit"],
            ["Hayden (Visibility CPA)", "Edit on co-maintained sheets only; View on others", "No access", "View", "View"],
            ["Cora service account", "View", "No access (PHI gating in code AND ACL)", "View", "View"],
            ["Anyone else", "No access", "No access", "No access", "No access"],
        ],
        col_widths=[1.5, 1.8, 1.8, 1.5, 1.0],
    )
    add_para(
        doc,
        "The Lex PHI tier exists because LBHS / Rita Tracking may contain client identifiers. "
        "Shaun (Lex Services) should sign off before any non-Lex eyes are given access. "
        "Until that's confirmed in writing, Hayden and Cora's service account both get no access "
        "to lex-lbhs_* sheets.",
    )

    add_hr(doc)

    # Anti-patterns
    doc.add_heading("Anti-patterns (don't do these)", level=2)
    anti_patterns = [
        ("Don't replace a _LIVE sheet with a fresh copy.", " Edit in place. The Google Sheet ID staying stable is the entire contract that lets Cora keep reading."),
        ("Don't put new files in visibility-binder/.", " That folder is legacy. New files go into the appropriate top-level folder per the structure above."),
        ("Don't let a Dropbox file sit overnight.", " Same-day pull. Otherwise the canonical source drifts back to Hayden's machine."),
        ("Don't trust an undated file.", " If you see something with no date prefix, rename it before saving - even if you have to guess at the period from the contents."),
        ("Don't share live-sheets/ links externally.", " These are canonical internal documents. If Hayden or anyone else needs a snapshot to send outside the org, export to a one-off file and share that."),
        ("Don't delete.", " Always archive into visibility-binder/_archive/ instead."),
    ]
    for bold_part, rest in anti_patterns:
        add_bullet_mixed(doc, [(bold_part, {"bold": True}), (rest, {})])

    add_hr(doc)

    # When unsure
    doc.add_heading("What to do when you're unsure", level=2)
    when_unsure = [
        ("Naming question?", " Read _shared/playbooks/naming-conventions.md first; if still unclear, ask Harrison."),
        ("\"Where does this file go?\"", " Look at the folder structure above. If nothing fits, ask Harrison before inventing a new folder."),
        ("\"Should I give Hayden access to X?\"", " Default to \"no\" until Harrison confirms. Permission grants are reversible but messy."),
        ("\"Cora is giving wrong numbers / not finding a file.\"", " Check (1) filename matches convention, (2) Cora's service account has view access, (3) file is in the expected folder. If all three pass, message Harrison."),
    ]
    for bold_part, rest in when_unsure:
        add_bullet_mixed(doc, [(bold_part, {"bold": True}), (rest, {})])

    add_hr(doc)

    # Related docs
    doc.add_heading("Related docs", level=2)
    related = [
        ("_shared/playbooks/naming-conventions.md", " - full naming rules (entity codes, files, tasks, scheduled tasks)"),
        ("_shared/playbooks/canonical-assets-doctrine.md", " - general principle: HJR-owned canonical copies, not third-party originals"),
        ("_shared/playbooks/cowork-storage-conventions.md", " - where Cowork puts files (Drive vs. local vs. OneDrive)"),
        ("01-HJR-Global/accounting/README.md", " - short pointer + at-a-glance reference for the accounting folder"),
    ]
    for path, desc in related:
        add_bullet_mixed(doc, [(path, {"mono": True}), (desc, {})])

    doc.save(str(SOP_OUT))
    print(f"Wrote {SOP_OUT}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    if not DRIVE_ROOT.exists():
        raise SystemExit(f"Output directory not found: {DRIVE_ROOT}\nIs Google Drive mounted?")
    build_readme()
    build_sop()
    print("\nDone. Attach these two files to Justin's email:")
    print(f"  - {README_OUT.name}")
    print(f"  - {SOP_OUT.name}")
