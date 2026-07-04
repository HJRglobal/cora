"""Drive sweep connector — full-corpus multi-user DWD ingestion.

For each account in monitored-email-accounts.yaml with drive_sweep: true,
impersonates the user via Domain-wide Delegation (DWD), enumerates all
Drive files modified within the freshness window, extracts text content
by MIME type, runs Claude Haiku classification to separate signal from
noise, then embeds and stores surviving chunks in the KB.

Auth requirements (one-time Harrison action):
  Add 'https://www.googleapis.com/auth/drive.readonly' to the Cora SA
  DWD grants in admin.google.com → Security → API Controls → Domain-wide
  Delegation → edit the SA entry → add scope.

Supported file types:
  Google Docs       → Drive export as text/plain
  Google Sheets     → Drive export as XLSX → openpyxl → Haiku summary
  Google Slides     → Drive export as text/plain
  PDFs              → download + pdfplumber text extraction
  XLSX/DOCX/TXT/CSV → download + best-available extractor

Deduplication:
  source='drive_sweep', source_id=file_id.  Same file_id shared to multiple
  users is only ingested once (first-seen wins via upsert).

Watermark:
  Stored via kb.set_sync_state(f"drive_sweep_{email}", iso_timestamp).
  Incremental re-runs only process files modified after last sweep.
"""

from __future__ import annotations

import io
import json
import logging
import os
import re
import textwrap
import time
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any

import yaml

from cora.connectors.drive_entity_detect import detect_entity_from_filename
from cora.drive_materializer import ENTITY_CODES as _MATERIALIZER_ENTITY_CODES
from cora.kb_exclusions import is_cora_internal_title
from cora.knowledge_base.store import Document
from cora.phi_guard import _PHI_PATTERNS

# Canonical parent entity codes (post sub-entity split), reused from the
# materializer so the two can't drift. O(1) membership for the W6-05 guard.
_CANONICAL_ENTITIES: frozenset[str] = frozenset(_MATERIALIZER_ENTITY_CODES)

log = logging.getLogger("cora.drive_sweep")

# ── MIME type routing ──────────────────────────────────────────────────────────

_GOOGLE_DOC_MIME = "application/vnd.google-apps.document"
_GOOGLE_SHEET_MIME = "application/vnd.google-apps.spreadsheet"
_GOOGLE_SLIDE_MIME = "application/vnd.google-apps.presentation"
_GOOGLE_FOLDER_MIME = "application/vnd.google-apps.folder"
_GOOGLE_FORM_MIME = "application/vnd.google-apps.form"
_GOOGLE_SCRIPT_MIME = "application/vnd.google-apps.script"

_SKIP_MIME_PREFIXES = (
    "image/",
    "video/",
    "audio/",
    "application/vnd.google-apps.folder",
    "application/vnd.google-apps.form",
    "application/vnd.google-apps.script",
    "application/vnd.google-apps.photo",
    "application/vnd.google-apps.drawing",
    "application/vnd.google-apps.map",
    "application/vnd.google-apps.shortcut",
    "application/zip",
    "application/gzip",
    "application/x-tar",
    "application/octet-stream",
)

_TEXT_MIME_TYPES = {
    "text/plain",
    "text/csv",
    "text/markdown",
    "text/html",
}

_PDF_MIME = "application/pdf"
_XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# Rows-per-tab ingested from any spreadsheet (Google Sheet export, XLSX binary,
# or Sheets-API fallback). Was 200, which silently truncated large data tables
# so later rows could never be recalled. Raised for data recall; the embedding
# layer now batches by token budget so large sheets no longer fail to embed.
_MAX_SHEET_ROWS = 5000

# ── Haiku classification ──────────────────────────────────────────────────────

_CLASSIFY_PROMPT = textwrap.dedent("""
You are classifying a business document for relevance to an HJR portfolio company.

File: {filename}
Owner: {user_name} ({user_email})
Likely entity: {entity}
Content preview:
---
{preview}
---

Rate business relevance 0-10:
  8-10 = contracts, signed agreements, financial data, meeting notes, strategic plans,
         operational SOPs, org charts, product specs, legal filings, cap tables,
         vendor agreements, staff rosters, compliance docs
  5-7  = internal communications, project notes, research, vendor correspondence
  2-4  = blank templates, rough drafts, duplicate copies, presentation shells
  0-1  = personal non-business files, temp files, system files, empty docs

Also suggest the most specific entity from:
  FNDR, HJRG, F3E, F3C, OSN, LEX, LEX-LLC, LEX-LTS, LEX-LBHS, LEX-LLA,
  UFL, BDM, HJRP, HJRP-CL, HJRP-LCI, HJRP-RR, HJRPROD

Respond with JSON only (no markdown):
{{"score": <0-10>, "entity": "<code>", "summary": "<one sentence>", "discard_reason": "<if score < 4, why>"}}
""").strip()


def _classify(anthropic_client: Any, filename: str, user_name: str,
              user_email: str, entity: str, preview: str) -> dict:
    """Run Haiku classification on a content preview. Returns parsed dict."""
    prompt = _CLASSIFY_PROMPT.format(
        filename=filename,
        user_name=user_name,
        user_email=user_email,
        entity=entity,
        preview=preview[:2000],
    )
    try:
        msg = anthropic_client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=200,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = msg.content[0].text.strip()
        # Strip markdown fences if Haiku adds them
        raw = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw, flags=re.DOTALL).strip()
        return json.loads(raw)
    except Exception as exc:
        log.warning("drive_sweep: Haiku classification failed for %s: %s", filename, exc)
        return {"score": 5, "entity": entity, "summary": filename, "discard_reason": ""}


# ── Content extraction helpers ────────────────────────────────────────────────

def _extract_google_doc(service: Any, file_id: str) -> str:
    """Export a Google Doc as plain text."""
    try:
        data = _retry_execute(service.files().export(
            fileId=file_id, mimeType="text/plain"
        ))
        return data.decode("utf-8", errors="replace") if isinstance(data, bytes) else data
    except Exception as exc:
        log.debug("drive_sweep: export failed for Doc %s: %s", file_id, exc)
        return ""


def _extract_google_sheet(service: Any, file_id: str,
                          sheets_service: Any | None = None) -> str:
    """Extract a Google Sheet's values as text.

    Primary path: Drive export as XLSX, parsed with openpyxl. Large sheets blow
    past Drive's export ceiling and raise HttpError 'exportSizeLimitExceeded',
    which previously dropped the file entirely. When that happens (or the export
    is empty) and a Sheets API service is available, fall back to the Sheets API
    `values` reader, which has no such size limit.
    """
    try:
        import openpyxl
    except ImportError:
        openpyxl = None  # type: ignore

    export_error: Exception | None = None
    if openpyxl is not None:
        try:
            data = _retry_execute(service.files().export(
                fileId=file_id,
                mimeType=_XLSX_MIME,
            ))
            wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
            parts: list[str] = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows_text: list[str] = []
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) for c in row if c is not None and str(c).strip()]
                    if cells:
                        rows_text.append("\t".join(cells))
                    if len(rows_text) >= _MAX_SHEET_ROWS:
                        break
                if rows_text:
                    parts.append(f"[Sheet: {sheet_name}]\n" + "\n".join(rows_text))
            text = "\n\n".join(parts)
            if text.strip():
                return text
        except Exception as exc:
            export_error = exc
            log.debug("drive_sweep: sheet export failed for %s: %s", file_id, exc)

    # Fallback: Sheets API values reader (no export size ceiling).
    if sheets_service is not None:
        text = _extract_sheet_via_api(sheets_service, file_id)
        if text.strip():
            if export_error is not None:
                log.info("drive_sweep: recovered oversized sheet %s via Sheets API", file_id)
            return text

    return ""


def _extract_sheet_via_api(sheets_service: Any, file_id: str) -> str:
    """Read a spreadsheet's values via the Sheets API (no export size limit).

    Lists each tab, then reads up to _MAX_SHEET_ROWS rows per tab. Cells are
    rendered as UNFORMATTED_VALUE so numbers/dates come through as data, not
    display strings. Returns the same `[Sheet: name]\\n<tab-joined rows>` text
    shape as the openpyxl path so downstream chunking is identical.
    """
    try:
        meta = _retry_execute(sheets_service.spreadsheets().get(
            spreadsheetId=file_id,
            fields="sheets.properties.title",
        ))
    except Exception as exc:
        log.debug("drive_sweep: Sheets API metadata failed for %s: %s", file_id, exc)
        return ""

    titles = [
        s.get("properties", {}).get("title", "")
        for s in meta.get("sheets", [])
        if s.get("properties", {}).get("title")
    ]
    if not titles:
        return ""

    parts: list[str] = []
    for title in titles:
        # Quote the tab title for the A1 range; escape embedded single quotes.
        safe_title = title.replace("'", "''")
        rng = f"'{safe_title}'!A1:ZZ{_MAX_SHEET_ROWS}"
        try:
            resp = _retry_execute(sheets_service.spreadsheets().values().get(
                spreadsheetId=file_id,
                range=rng,
                valueRenderOption="UNFORMATTED_VALUE",
            ))
        except Exception as exc:
            log.debug("drive_sweep: Sheets API values failed for %s/%s: %s",
                      file_id, title, exc)
            continue
        rows = resp.get("values", [])
        rows_text: list[str] = []
        for row in rows:
            cells = [str(c) for c in row if c is not None and str(c).strip()]
            if cells:
                rows_text.append("\t".join(cells))
        if rows_text:
            parts.append(f"[Sheet: {title}]\n" + "\n".join(rows_text))
    return "\n\n".join(parts)


def _extract_pdf_bytes(data: bytes) -> str:
    """Extract text from PDF bytes using pdfplumber."""
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            pages: list[str] = []
            for page in pdf.pages[:40]:  # cap at 40 pages
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n\n".join(pages)
    except ImportError:
        log.warning("drive_sweep: pdfplumber not installed -- skipping PDF")
        return ""
    except Exception as exc:
        log.debug("drive_sweep: pdfplumber extraction failed: %s", exc)
        return ""


def _extract_xlsx_bytes(data: bytes) -> str:
    """Parse XLSX binary with openpyxl."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
        parts: list[str] = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows_text: list[str] = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None and str(c).strip()]
                if cells:
                    rows_text.append("\t".join(cells))
                if len(rows_text) >= _MAX_SHEET_ROWS:
                    break
            if rows_text:
                parts.append(f"[Sheet: {sheet_name}]\n" + "\n".join(rows_text))
        return "\n\n".join(parts)
    except Exception as exc:
        log.debug("drive_sweep: XLSX parse failed: %s", exc)
        return ""


def _download_and_extract(service: Any, file_meta: dict) -> str:
    """Download a binary file and extract text based on MIME type."""
    mime = file_meta.get("mimeType", "")
    file_id = file_meta["id"]
    try:
        data: bytes = _retry_execute(service.files().get_media(fileId=file_id))
    except Exception as exc:
        log.debug("drive_sweep: download failed for %s: %s", file_id, exc)
        return ""

    if mime == _PDF_MIME:
        return _extract_pdf_bytes(data)
    if mime == _XLSX_MIME:
        return _extract_xlsx_bytes(data)
    if mime in _TEXT_MIME_TYPES or mime.startswith("text/"):
        return data.decode("utf-8", errors="replace")
    if mime == _DOCX_MIME:
        # Try python-docx if available, otherwise treat as unsupported
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(data))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            pass
    return ""


def _extract_content(service: Any, file_meta: dict,
                     sheets_service: Any | None = None) -> str:
    """Route to the right extractor for a Drive file."""
    mime = file_meta.get("mimeType", "")
    file_id = file_meta["id"]

    # Skip known-noise MIME types immediately
    for prefix in _SKIP_MIME_PREFIXES:
        if mime.startswith(prefix) or mime == prefix:
            return ""

    if mime == _GOOGLE_DOC_MIME:
        return _extract_google_doc(service, file_id)
    if mime == _GOOGLE_SHEET_MIME:
        return _extract_google_sheet(service, file_id, sheets_service=sheets_service)
    if mime in (_GOOGLE_SLIDE_MIME,):
        return _extract_google_doc(service, file_id)  # same export path
    # Binary files
    return _download_and_extract(service, file_meta)


# ── Chunking + KB storage ─────────────────────────────────────────────────────

_CHUNK_SIZE = 1400   # characters per KB chunk
_CHUNK_OVERLAP = 150


def _chunk_text(text: str) -> list[str]:
    """Split long text into overlapping chunks."""
    if len(text) <= _CHUNK_SIZE:
        return [text]
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = start + _CHUNK_SIZE
        chunks.append(text[start:end])
        start = end - _CHUNK_OVERLAP
    return chunks


def _ingest_file(kb: Any, file_meta: dict, content: str,
                 classification: dict, user: dict) -> int:
    """Store classified content in the KB as a single Document. Returns chunk count ingested."""
    entity = classification.get("entity") or user.get("entity_default", "FNDR")
    sub_entity: str | None = None
    if "-" in entity:
        prefix = entity.split("-")[0]
        if prefix in ("LEX", "HJRP", "HJRPROD"):
            sub_entity = entity
            entity = prefix

    # Entity-firewall guard (audit W6-05): Haiku can hallucinate an off-menu code
    # (e.g. the filename token "F3" minted entity='F3' from an OSN receipt) that
    # no channel routes to. Reject any non-canonical post-split parent code and
    # fall back to the file owner's canonical default rather than minting a novel
    # entity. _CANONICAL_ENTITIES is drive_materializer.ENTITY_CODES (the 10 parent
    # codes; the classifier prompt's allowed set collapses to exactly these after
    # the sub-entity split above). Mirrors session_capture's VALID_ENTITIES guard.
    if entity not in _CANONICAL_ENTITIES:
        default = user.get("entity_default", "FNDR")
        fallback = default if default in _CANONICAL_ENTITIES else "FNDR"
        log.warning(
            "drive_sweep: non-canonical entity %r for %r -> falling back to %s",
            entity, file_meta.get("name"), fallback,
        )
        entity = fallback
        sub_entity = None

    summary = classification.get("summary", file_meta["name"])
    file_id = file_meta["id"]
    modified_iso = file_meta.get("modifiedTime", "")
    drive_link = f"https://drive.google.com/file/d/{file_id}/view"

    date_modified: int | None = None
    if modified_iso:
        try:
            date_modified = int(datetime.fromisoformat(modified_iso.replace("Z", "+00:00")).timestamp())
        except Exception:
            pass

    doc = Document(
        source="drive_sweep",
        source_id=file_id,
        entity=entity,
        sub_entity=sub_entity,
        title=file_meta["name"],
        content=content,
        deep_link=drive_link,
        author=user.get("name", user["email"]),
        date_modified=date_modified,
        metadata={
            "mime_type": file_meta.get("mimeType", ""),
            "user_email": user["email"],
            "user_name": user.get("name", user["email"]),
            "haiku_score": classification.get("score", 5),
            "haiku_summary": summary,
            "modified_time": modified_iso,
            "drive_link": drive_link,
        },
    )
    try:
        return kb.upsert_documents([doc])
    except Exception as exc:
        log.warning("drive_sweep: upsert_documents failed for %s: %s", file_meta["name"], exc)
        return 0


# ── Per-user sweep ────────────────────────────────────────────────────────────

_DRIVE_SCOPE = "https://www.googleapis.com/auth/drive.readonly"
_SHEETS_SCOPE = "https://www.googleapis.com/auth/spreadsheets.readonly"

_SUPPORTED_MIME_QUERY = " or ".join([
    f"mimeType='{_GOOGLE_DOC_MIME}'",
    f"mimeType='{_GOOGLE_SHEET_MIME}'",
    f"mimeType='{_GOOGLE_SLIDE_MIME}'",
    f"mimeType='{_PDF_MIME}'",
    f"mimeType='{_XLSX_MIME}'",
    f"mimeType='{_DOCX_MIME}'",
    "mimeType='text/plain'",
    "mimeType='text/csv'",
    "mimeType='text/markdown'",
])

# Files from these folders won't get an entity override (they're already labelled)
_NOISE_FOLDER_PATTERNS = re.compile(
    r"(trash|archive|old|backup|template|\\.cache|\\.tmp|receipts?/personal)",
    re.IGNORECASE,
)


def _retry_execute(request: Any, max_retries: int = 3) -> Any:
    """Execute a Google API request, retrying on 429 (rate limit) or 503 (unavailable).

    Sleeps exponentially (2s, 4s, 8s) between retries. Raises the original
    HttpError if all retries are exhausted.
    """
    from googleapiclient.errors import HttpError

    for attempt in range(max_retries + 1):
        try:
            return request.execute()
        except HttpError as exc:
            status = exc.resp.status if exc.resp else 0
            if status in (429, 503) and attempt < max_retries:
                sleep_seconds = 2 ** (attempt + 1)
                log.warning(
                    "drive_sweep: Google API returned %d — retrying in %ds (attempt %d/%d)",
                    status, sleep_seconds, attempt + 1, max_retries,
                )
                time.sleep(sleep_seconds)
                continue
            raise


def _build_drive_service(sa_json_path: str, user_email: str):
    """Build a Drive v3 service impersonating user_email via DWD."""
    from google.oauth2 import service_account
    from googleapiclient.discovery import build

    creds = service_account.Credentials.from_service_account_file(
        sa_json_path, scopes=[_DRIVE_SCOPE]
    ).with_subject(user_email)
    return build("drive", "v3", credentials=creds, cache_discovery=False)


def _build_sheets_service(sa_json_path: str, user_email: str):
    """Build a Sheets v4 service impersonating user_email via DWD.

    Requires the Cora SA DWD grant to include the scope
    'https://www.googleapis.com/auth/spreadsheets.readonly' (one-time Harrison
    action in admin.google.com). Returns None if the service cannot be built so
    the caller degrades gracefully to the Drive-export path.
    """
    try:
        from google.oauth2 import service_account
        from googleapiclient.discovery import build

        creds = service_account.Credentials.from_service_account_file(
            sa_json_path, scopes=[_SHEETS_SCOPE]
        ).with_subject(user_email)
        return build("sheets", "v4", credentials=creds, cache_discovery=False)
    except Exception as exc:
        log.warning("drive_sweep: could not build Sheets service for %s: %s",
                    user_email, exc)
        return None


def sweep_user(
    user: dict,
    sa_json_path: str,
    kb: Any,
    anthropic_client: Any,
    freshness_days: int = 730,
    dry_run: bool = False,
    seen_file_ids: set[str] | None = None,
) -> dict:
    """Sweep one user's Drive. Returns stats dict."""
    email = user["email"]
    name = user.get("name", email)
    entity_default = user.get("entity_default", "FNDR")
    is_lex_user = email.endswith("@lexingtonservices.com")

    if seen_file_ids is None:
        seen_file_ids = set()

    stats = {"files_enumerated": 0, "files_extracted": 0,
             "chunks_ingested": 0, "phi_skipped": 0, "noise_filtered": 0,
             "dedup_skipped": 0}

    # Watermark for incremental sync
    watermark_key = f"drive_sweep_{email}"
    checkpoint_key = f"drive_sweep_checkpoint_{email}"
    last_run_ts: int | None = None
    try:
        state = kb.get_sync_state(watermark_key)
        if state and isinstance(state[0], int):
            last_run_ts = state[0]
    except Exception:
        pass

    # Resume checkpoint: if a previous sweep was interrupted mid-user, pick up
    # from the last saved nextPageToken rather than starting from scratch.
    resume_page_token: str | None = None
    files_processed_base: int = 0
    if not dry_run:
        try:
            ckpt = kb.get_checkpoint(checkpoint_key)
            if ckpt:
                resume_page_token = ckpt.get("page_token")
                files_processed_base = int(ckpt.get("files_processed", 0))
                log.info(
                    "drive_sweep: resuming %s from checkpoint "
                    "(page_token=%s files_processed=%d)",
                    email,
                    resume_page_token or "(first page)",
                    files_processed_base,
                )
        except Exception:
            pass

    cutoff = datetime.now(timezone.utc) - timedelta(days=freshness_days)
    if last_run_ts and not resume_page_token:
        # Only apply watermark cutoff on fresh (non-resume) runs so the query
        # matches the same result set that was being paged when interrupted.
        watermark_dt = datetime.fromtimestamp(last_run_ts, tz=timezone.utc)
        cutoff = max(cutoff, watermark_dt)

    cutoff_str = cutoff.strftime("%Y-%m-%dT%H:%M:%S")
    log.info("drive_sweep: sweeping %s (%s) modified since %s", email, name, cutoff_str)

    try:
        service = _build_drive_service(sa_json_path, email)
    except Exception as exc:
        log.error("drive_sweep: could not build Drive service for %s: %s", email, exc)
        return stats

    # Sheets API service for the oversized-export fallback (None if scope/grant
    # missing -> we degrade to the Drive-export path without crashing).
    sheets_service = _build_sheets_service(sa_json_path, email)

    # Build Drive files.list query
    q = (
        f"({_SUPPORTED_MIME_QUERY})"
        f" and trashed = false"
        f" and modifiedTime > '{cutoff_str}'"
    )

    page_token: str | None = resume_page_token
    files_processed_this_run: int = files_processed_base
    run_start = datetime.now(timezone.utc)

    while True:
        try:
            list_kwargs: dict = dict(
                q=q,
                spaces="drive",
                fields="nextPageToken, files(id, name, mimeType, modifiedTime, size, webViewLink, parents)",
                pageSize=100,
            )
            if page_token:
                list_kwargs["pageToken"] = page_token
            resp = service.files().list(**list_kwargs).execute()
        except Exception as exc:
            log.error("drive_sweep: files.list failed for %s: %s", email, exc)
            break

        files = resp.get("files", [])
        stats["files_enumerated"] += len(files)

        for file_meta in files:
            file_id = file_meta["id"]
            filename = file_meta.get("name", "")

            # Cross-user dedup
            if file_id in seen_file_ids:
                stats["dedup_skipped"] += 1
                continue
            seen_file_ids.add(file_id)

            # WS1: never ingest Cora's OWN build/audit/forensic docs or runtime logs.
            # drive_sweep walks the Founder OS Drive tree, so these would otherwise
            # land in the KB under a Drive-file-id source_id no path rule can catch,
            # and Cora would recite her own audit notes as fact (the Minute Press miss).
            # broad=True: fail-safe to the WIDER exclusion at ingest -- over-excluding
            # Cora's own ops docs from the KB is harmless; under-excluding re-opens the
            # self-diagnostic leak (the targeted set alone misses cora-code-* etc.).
            if is_cora_internal_title(filename, broad=True):
                stats.setdefault("cora_internal_skipped", 0)
                stats["cora_internal_skipped"] += 1
                continue

            # Skip very small files (likely empty/template)
            size_str = file_meta.get("size", "0")
            try:
                if int(size_str) < 200:
                    stats["noise_filtered"] += 1
                    continue
            except (ValueError, TypeError):
                pass

            # Extract content
            content = _extract_content(service, file_meta, sheets_service=sheets_service)
            if not content or len(content.strip()) < 150:
                stats["noise_filtered"] += 1
                continue

            # PHI guard for Lex users — quarantine before classification
            if is_lex_user and _PHI_PATTERNS.search(content[:5000]):
                log.debug("drive_sweep: PHI guard triggered for %s/%s", email, filename)
                stats["phi_skipped"] += 1
                continue

            stats["files_extracted"] += 1
            files_processed_this_run += 1

            # Haiku classification
            preview = content[:2000]
            classification = _classify(
                anthropic_client, filename, name, email, entity_default, preview
            )
            score = classification.get("score", 0)

            # Deterministic filename entity override -- HJR naming convention is
            # more reliable than Haiku for entity attribution (Haiku has tagged
            # OSN P&Ls and HJRP invoices as LEX). Score still comes from Haiku.
            detected_entity = detect_entity_from_filename(filename)
            if detected_entity:
                classification["entity"] = detected_entity

            if score < 4:
                log.debug("drive_sweep: discarded %s (score=%s: %s)",
                          filename, score, classification.get("discard_reason", ""))
                stats["noise_filtered"] += 1
                continue

            log.info("drive_sweep: ingesting %s (score=%s, entity=%s)",
                     filename, score, classification.get("entity", entity_default))

            if dry_run:
                stats["chunks_ingested"] += 1
                continue

            n = _ingest_file(kb, file_meta, content, classification, user)
            stats["chunks_ingested"] += n

        page_token = resp.get("nextPageToken")

        # Persist checkpoint after each page so a killed/crashed sweep can resume
        # from here next run instead of restarting from page 1.
        if not dry_run:
            try:
                kb.set_checkpoint(checkpoint_key, {
                    "page_token": page_token,           # None when this was the last page
                    "files_processed": files_processed_this_run,
                })
            except Exception as exc:
                log.warning("drive_sweep: could not save checkpoint for %s: %s", email, exc)

        if not page_token:
            break

    # Advance watermark only on successful (non-dry-run) sweep
    if not dry_run:
        try:
            kb.set_sync_state(watermark_key, int(run_start.timestamp()))
        except Exception as exc:
            log.warning("drive_sweep: could not advance watermark for %s: %s", email, exc)
        # Sweep completed — clear the checkpoint so the next run starts fresh
        try:
            kb.delete_checkpoint(checkpoint_key)
        except Exception as exc:
            log.warning("drive_sweep: could not clear checkpoint for %s: %s", email, exc)

    log.info(
        "drive_sweep: %s done -- enumerated=%d extracted=%d ingested=%d "
        "phi_skipped=%d noise=%d dedup=%d",
        email,
        stats["files_enumerated"], stats["files_extracted"], stats["chunks_ingested"],
        stats["phi_skipped"], stats["noise_filtered"], stats["dedup_skipped"],
    )
    return stats


# ── Main entry point ──────────────────────────────────────────────────────────

def run_sweep(
    sa_json_path: str,
    accounts_yaml_path: str,
    kb: Any,
    anthropic_client: Any,
    freshness_days: int = 730,
    dry_run: bool = False,
    only_email: str | None = None,
) -> dict:
    """Sweep all enabled drive_sweep accounts. Returns aggregate stats."""
    with open(accounts_yaml_path, encoding="utf-8") as f:
        cfg = yaml.safe_load(f)

    accounts = [
        a for a in cfg.get("accounts", [])
        if a.get("enabled", True) and a.get("dwd_eligible", False)
        and a.get("drive_sweep", False)
    ]

    if only_email:
        accounts = [a for a in accounts if a["email"] == only_email]

    log.info("drive_sweep: starting sweep of %d accounts (freshness=%dd, dry_run=%s)",
             len(accounts), freshness_days, dry_run)

    aggregate: dict = {
        "accounts_swept": 0, "files_enumerated": 0, "files_extracted": 0,
        "chunks_ingested": 0, "phi_skipped": 0, "noise_filtered": 0,
        "dedup_skipped": 0,
    }
    seen_file_ids: set[str] = set()

    for user in accounts:
        stats = sweep_user(
            user=user,
            sa_json_path=sa_json_path,
            kb=kb,
            anthropic_client=anthropic_client,
            freshness_days=freshness_days,
            dry_run=dry_run,
            seen_file_ids=seen_file_ids,
        )
        aggregate["accounts_swept"] += 1
        for k in ("files_enumerated", "files_extracted", "chunks_ingested",
                  "phi_skipped", "noise_filtered", "dedup_skipped"):
            aggregate[k] += stats.get(k, 0)

    log.info(
        "drive_sweep: COMPLETE -- accounts=%d enumerated=%d extracted=%d "
        "ingested=%d phi_skipped=%d noise=%d dedup=%d",
        aggregate["accounts_swept"], aggregate["files_enumerated"],
        aggregate["files_extracted"], aggregate["chunks_ingested"],
        aggregate["phi_skipped"], aggregate["noise_filtered"],
        aggregate["dedup_skipped"],
    )
    return aggregate


# ── HJR-Founder-OS shared folder sweep ────────────────────────────────────────

FOUNDERS_OS_ROOT_ID = "1TfxuKxzXz0-NipAFYqbK5AxowAy-LIPG"

_FOUNDERS_OS_ENTITY_MAP: dict[str, str] = {
    "00-founder":            "FNDR",
    "01-hjr-global":         "HJRG",
    "02-f3-energy":          "F3E",
    "03-f3-community":       "F3C",
    "04-ufl":                "UFL",
    "05-hjr-productions":    "HJRPROD",
    "06-hjr-properties":     "HJRP",
    "07-big-d-media":        "BDM",
    "08-lexington-services": "LEX",
    "09-one-stop-nutrition":  "OSN",
    "_shared":               "FNDR",
}

_HJRP_SUB_ENTITY: dict[str, str] = {
    "cinema-lanes": "HJRP-CL",
    "lci-realty":   "HJRP-LCI",
    "rogers-ranch": "HJRP-RR",
}

_LEX_SUB_ENTITY: dict[str, str] = {
    "llc":               "LEX-LLC",
    "lexington-llc":     "LEX-LLC",
    "lts":               "LEX-LTS",
    "lexington-therapy": "LEX-LTS",
    "lbhs":              "LEX-LBHS",
    "behavioral-health": "LEX-LBHS",
    "lla":               "LEX-LLA",
    "lex-life-academy":  "LEX-LLA",
    "lex-life":          "LEX-LLA",
}

_ENTITY_SUB_MAP: dict[str, dict[str, str]] = {
    "LEX":  _LEX_SUB_ENTITY,
    "HJRP": _HJRP_SUB_ENTITY,
}

_FOUNDERS_OS_SKIP_FOLDERS = frozenset({
    "_archive", "archive", "old", "backup", "temp", ".cache",
    ".tmp", "node_modules", ".git",
    # Drive-materialization (2026-06-29): the nightly _brain/swept/ digests. _brain is
    # already unmapped here (no entity -> skipped), so this is defense-in-depth in case
    # _brain is ever mapped; it must NEVER feed back into the KB.
    "swept",
})

_LEX_SCORE_THRESHOLD = 6
_DEFAULT_SCORE_THRESHOLD = 4


def _build_sa_drive_service_direct(sa_json_path: str) -> Any:
    """Build a Drive v3 service using direct SA credentials (no DWD impersonation).

    Used for the HJR-Founder-OS shared folder. The SA must have the folder
    shared with it as Viewer in Google Drive (one-time Harrison action).
    SA email: cora-calendar@cora-calendar-readonly.iam.gserviceaccount.com
    """
    from google.oauth2 import service_account
    from googleapiclient.discovery import build

    creds = service_account.Credentials.from_service_account_file(
        sa_json_path, scopes=[_DRIVE_SCOPE]
    )
    return build("drive", "v3", credentials=creds, cache_discovery=False)


def _build_sa_sheets_service_direct(sa_json_path: str) -> Any:
    """Build a Sheets v4 service using direct SA credentials (no impersonation).

    Used for the HJR-Founder-OS shared folder oversized-sheet fallback. Returns
    None on any failure so the caller degrades to the Drive-export path.
    """
    try:
        from google.oauth2 import service_account
        from googleapiclient.discovery import build

        creds = service_account.Credentials.from_service_account_file(
            sa_json_path, scopes=[_SHEETS_SCOPE]
        )
        return build("sheets", "v4", credentials=creds, cache_discovery=False)
    except Exception as exc:
        log.warning("founders_os: could not build Sheets service: %s", exc)
        return None


def _list_subfolders(service: Any, parent_id: str) -> list[dict]:
    """List immediate subfolder children of a Drive folder."""
    results: list[dict] = []
    page_token: str | None = None
    while True:
        kwargs: dict = dict(
            q=(
                f"'{parent_id}' in parents"
                f" and mimeType='{_GOOGLE_FOLDER_MIME}'"
                f" and trashed=false"
            ),
            fields="nextPageToken, files(id, name)",
            pageSize=100,
        )
        if page_token:
            kwargs["pageToken"] = page_token
        try:
            resp = _retry_execute(service.files().list(**kwargs))
        except Exception as exc:
            log.warning("founders_os: subfolder list failed for %s: %s", parent_id, exc)
            break
        results.extend(resp.get("files", []))
        page_token = resp.get("nextPageToken")
        if not page_token:
            break
    return results


def _process_single_folder_files(
    *,
    service: Any,
    folder_id: str,
    label: str,
    effective_entity: str,
    kb: Any,
    anthropic_client: Any,
    cutoff_str: str,
    dry_run: bool,
    is_lex: bool,
    score_threshold: int,
    seen_file_ids: set,
    stats: dict,
    sheets_service: Any | None = None,
    deadline_monotonic: float | None = None,
) -> bool:
    """Process files directly inside one folder (non-recursive, 'in parents' only).

    Returns True if the folder was fully drained (all pages processed), False if
    it stopped early because the wall-clock budget (deadline_monotonic) elapsed.
    W4-01: the caller must only mark a folder "done" in its checkpoint when this
    returns True, so a budget-interrupted folder is re-processed next run (an
    idempotent upsert on source_id=file_id — re-work, never data loss)."""
    q = (
        f"'{folder_id}' in parents"
        f" and ({_SUPPORTED_MIME_QUERY})"
        f" and trashed = false"
        f" and modifiedTime > '{cutoff_str}'"
    )
    page_token: str | None = None
    while True:
        list_kwargs: dict = dict(
            q=q,
            spaces="drive",
            fields=(
                "nextPageToken,"
                "files(id,name,mimeType,modifiedTime,size,webViewLink,parents)"
            ),
            pageSize=100,
        )
        if page_token:
            list_kwargs["pageToken"] = page_token

        try:
            resp = _retry_execute(service.files().list(**list_kwargs))
        except Exception as exc:
            log.error("founders_os: files.list failed for %s/%s: %s", label, folder_id, exc)
            break

        files = resp.get("files", [])
        stats["files_enumerated"] += len(files)

        for file_meta in files:
            file_id  = file_meta["id"]
            filename = file_meta.get("name", "")

            if file_id in seen_file_ids:
                stats["dedup_skipped"] += 1
                continue
            seen_file_ids.add(file_id)

            # WS1: never ingest Cora's OWN build/audit/forensic docs or runtime logs.
            # broad=True: fail-safe to the WIDER exclusion at ingest -- over-excluding
            # Cora's own ops docs from the KB is harmless; under-excluding re-opens the
            # self-diagnostic leak (the targeted set alone misses cora-code-* etc.).
            if is_cora_internal_title(filename, broad=True):
                stats.setdefault("cora_internal_skipped", 0)
                stats["cora_internal_skipped"] += 1
                continue

            try:
                if int(file_meta.get("size", "0")) < 200:
                    stats["noise_filtered"] += 1
                    continue
            except (ValueError, TypeError):
                pass

            content = _extract_content(service, file_meta, sheets_service=sheets_service)
            if not content or len(content.strip()) < 150:
                stats["noise_filtered"] += 1
                continue

            if is_lex and _PHI_PATTERNS.search(content[:5000]):
                log.debug("founders_os: PHI guard triggered for %s / %s", label, filename)
                stats["phi_skipped"] += 1
                continue

            stats["files_extracted"] += 1

            classification = _classify(
                anthropic_client, filename,
                "HJR-Founder-OS", "founders_os@hjrglobal.com",
                effective_entity, content[:2000],
            )
            score = classification.get("score", 0)
            classification["entity"] = effective_entity  # folder path wins over Haiku

            if score < score_threshold:
                log.debug("founders_os: discarded %s (score=%s)", filename, score)
                stats["noise_filtered"] += 1
                continue

            log.info("founders_os: ingesting %s / %s (score=%s)", label, filename, score)

            if dry_run:
                stats["chunks_ingested"] += 1
                log.info("[DRY RUN] Would ingest: %s -> entity=%s", filename, effective_entity)
                continue

            user_dict = {
                "email": "founders_os@hjrglobal.com",
                "name":  "HJR-Founder-OS",
                "entity_default": effective_entity,
            }
            n = _ingest_file(kb, file_meta, content, classification, user_dict)
            stats["chunks_ingested"] += n

        page_token = resp.get("nextPageToken")
        if not page_token:
            break

        # W4-01: yield between pages if the wall-clock budget elapsed, so a folder
        # with thousands of files can't blow past the task's ExecutionTimeLimit
        # (which SIGKILLs mid-commit). Returning False signals "not fully drained"
        # -> the caller leaves this folder OUT of its completed set, so next run
        # re-processes it from page 1 (idempotent upsert; bounded to one folder).
        if deadline_monotonic is not None and time.monotonic() >= deadline_monotonic:
            return False

    return True


def _sweep_folder_tree(
    *,
    service: Any,
    folder_id: str,
    entity: str,
    sub_entity: str | None,
    kb: Any,
    anthropic_client: Any,
    cutoff_str: str,
    dry_run: bool,
    is_lex: bool,
    score_threshold: int,
    seen_file_ids: set,
    stats: dict,
    checkpoint_key: str,
    sheets_service: Any | None = None,
    deadline_monotonic: float | None = None,
    skip_folder_ids: frozenset[str] | None = None,
) -> bool:
    """BFS walk of a folder subtree using 'in parents' queries at each depth.

    The Drive API 'in ancestors' operator only works for Shared Drives, not
    personal My Drive folders. This BFS approach works for all Drive types:
    enumerate files in each folder with 'in parents', then recurse into subfolders.

    Entity is pre-determined from folder context; Haiku only scores relevance.
    PHI guard fires for all LEX content before classification.
    Dedup via seen_file_ids prevents double-ingesting files in overlapping sweeps.

    W4-01 resumable sweep: ``checkpoint_key`` is now consumed. Each folder whose
    files are fully processed is recorded in a per-subtree checkpoint
    (``completed_folder_ids``); a subtree that finishes all folders records
    ``tree_done: True``. On a resumed run the checkpoint lets the walk skip the
    file-processing of already-completed folders (it still re-lists them to
    rebuild the frontier — cheap) so a subtree bigger than one budget window
    (LEX) chips away across runs instead of restarting each day. Returns True if
    the whole subtree completed, False if the wall-clock budget elapsed first;
    the caller advances the entity watermark ONLY when every subtree returns True,
    so a budget interruption can never half-advance a watermark.
    """
    effective_entity = sub_entity or entity
    label = f"{entity}/{sub_entity}" if sub_entity else entity
    skip_folder_ids = skip_folder_ids or frozenset()

    # Resume state: load which folders in this subtree are already done.
    completed: set[str] = set()
    if not dry_run:
        try:
            ckpt = kb.get_checkpoint(checkpoint_key)
            if ckpt:
                if ckpt.get("tree_done"):
                    # Whole subtree finished in a prior run — nothing to redo.
                    return True
                completed = {str(x) for x in ckpt.get("completed_folder_ids", [])}
                if completed:
                    log.info(
                        "founders_os: resuming %s subtree from checkpoint "
                        "(%d folders already done)", label, len(completed),
                    )
        except Exception as exc:  # noqa: BLE001 — a bad checkpoint never blocks a sweep
            log.warning("founders_os: checkpoint read failed for %s: %s", label, exc)

    def _persist(tree_done: bool) -> None:
        if dry_run:
            return
        try:
            kb.set_checkpoint(checkpoint_key, {
                "completed_folder_ids": sorted(completed),
                "tree_done": tree_done,
            })
        except Exception as exc:  # noqa: BLE001
            log.warning("founders_os: checkpoint save failed for %s: %s", label, exc)

    # BFS queue: process one folder at a time, recurse into subfolders
    queue: list[str] = [folder_id]
    visited: set[str] = set()

    while queue:
        current_id = queue.pop(0)
        if current_id in visited or current_id in skip_folder_ids:
            continue
        visited.add(current_id)

        # Skip file-processing of folders already completed in a prior run, but
        # still discover their subfolders below so the frontier is rebuilt.
        if current_id not in completed:
            folder_done = _process_single_folder_files(
                service=service,
                folder_id=current_id,
                label=label,
                effective_entity=effective_entity,
                kb=kb,
                anthropic_client=anthropic_client,
                cutoff_str=cutoff_str,
                dry_run=dry_run,
                is_lex=is_lex,
                score_threshold=score_threshold,
                seen_file_ids=seen_file_ids,
                stats=stats,
                sheets_service=sheets_service,
                deadline_monotonic=deadline_monotonic,
            )
            if not folder_done:
                # Budget elapsed mid-folder: do NOT mark this folder done.
                _persist(tree_done=False)
                return False
            completed.add(current_id)
            _persist(tree_done=False)

        # Discover subfolders and add to BFS queue
        subfolders = _list_subfolders(service, current_id)
        for subfolder in subfolders:
            sub_name = subfolder["name"].lower()
            if (sub_name not in _FOUNDERS_OS_SKIP_FOLDERS
                    and subfolder["id"] not in visited
                    and subfolder["id"] not in skip_folder_ids):
                queue.append(subfolder["id"])

        # Budget check between folders — stop cleanly so the checkpoint above is
        # the resume point next run (never a SIGKILL mid-commit).
        if deadline_monotonic is not None and time.monotonic() >= deadline_monotonic:
            _persist(tree_done=False)
            return False

    # Whole subtree walked — record it so a resumed run skips it instantly.
    _persist(tree_done=True)
    return True


def _founders_os_entity_for(folder_name: str) -> str | None:
    """Map a top-level folder name to its entity code (exact, then prefix)."""
    folder_key = folder_name.lower()
    if folder_key in _FOUNDERS_OS_SKIP_FOLDERS:
        return None
    entity = _FOUNDERS_OS_ENTITY_MAP.get(folder_key)
    if entity is None:
        for key, ent in _FOUNDERS_OS_ENTITY_MAP.items():
            if folder_key.startswith(key):
                return ent
    return entity


def sweep_founders_os(
    sa_json_path: str,
    kb: Any,
    anthropic_client: Any,
    root_folder_id: str = FOUNDERS_OS_ROOT_ID,
    entity_filter: str | None = None,
    freshness_days: int = 730,
    dry_run: bool = False,
    time_budget_min: int | None = None,
) -> dict:
    """Sweep the HJR-Founder-OS shared Drive folder into Cora's KB.

    Entity is determined from folder path (deterministic), not Haiku.
    Haiku only scores relevance (0-10); files below threshold are discarded.
    PHI guard enforced on all LEX content. LEX threshold = 6, others = 4.

    W4-01 reliability: ``time_budget_min`` bounds the wall clock. Before this
    fix the scheduled task (PT2H ExecutionTimeLimit) SIGKILLed the process
    mid-LEX-ingest EVERY run — so 7 of 10 entities never completed a sweep and
    LEX re-scanned from scratch daily. Now the sweep (1) processes the neediest
    entities first (never-completed, then stalest-watermark), (2) resumes big
    subtrees from a checkpoint instead of restarting, and (3) stops CLEANLY a
    few minutes before the task limit, persisting progress, so a subtree larger
    than one window converges over ≤ a couple of runs. A watermark advances only
    when EVERY subtree of an entity completed, so a budget cut can never
    half-advance one. ``time_budget_min`` None/<=0 = unlimited (manual backfill).

    One-time Harrison action required: share HJR-Founder-OS folder with
    cora-calendar@cora-calendar-readonly.iam.gserviceaccount.com as Viewer.
    """
    log.info(
        "founders_os: starting sweep root=%s filter=%s freshness=%dd dry_run=%s budget=%s",
        root_folder_id, entity_filter or "ALL", freshness_days, dry_run,
        f"{time_budget_min}min" if time_budget_min and time_budget_min > 0 else "none",
    )

    try:
        service = _build_sa_drive_service_direct(sa_json_path)
    except Exception as exc:
        log.error("founders_os: could not build Drive service: %s", exc)
        return {"error": str(exc)}

    # Sheets API service for oversized-export fallback (None -> degrade gracefully).
    sheets_service = _build_sa_sheets_service_direct(sa_json_path)

    allowed_entities: set[str] | None = None
    if entity_filter:
        allowed_entities = {e.strip().upper() for e in entity_filter.split(",")}

    # Wall-clock budget: honoured only on live runs (a dry-run preview should
    # never be cut off, and it writes no checkpoint to resume from anyway).
    deadline: float | None = None
    if not dry_run and time_budget_min and time_budget_min > 0:
        deadline = time.monotonic() + time_budget_min * 60

    top_folders = _list_subfolders(service, root_folder_id)
    log.info("founders_os: found %d top-level folders", len(top_folders))

    # ── Build + order the entity work-list (W4-01 neediest-first) ───────────────
    # Never-completed entities (no watermark row) sort first, then stalest
    # watermark first; fresh incremental entities go last. This keeps the entity
    # the SIGKILL used to starve (LEX / anything after it in Drive's arbitrary
    # folder order) from being perpetually skipped, and lets it resume + finish
    # before the cheap incremental entities consume any remaining budget.
    work: list[dict] = []
    for folder in top_folders:
        entity = _founders_os_entity_for(folder["name"])
        if entity is None:
            log.info("founders_os: no mapping for %r — skipping", folder["name"])
            continue
        if allowed_entities and entity not in allowed_entities:
            continue
        watermark_key = f"founders_os_{entity}_{folder['id']}"
        watermark_ts: int | None = None
        try:
            state = kb.get_sync_state(watermark_key)
            if state and isinstance(state[0], int):
                watermark_ts = state[0]
        except Exception:
            pass
        work.append({
            "folder_id": folder["id"], "folder_name": folder["name"],
            "entity": entity, "watermark_key": watermark_key,
            "watermark_ts": watermark_ts,
        })
    work.sort(key=lambda w: (w["watermark_ts"] is not None, w["watermark_ts"] or 0))

    aggregate: dict = {
        "entities_swept": 0, "files_enumerated": 0, "files_extracted": 0,
        "chunks_ingested": 0, "phi_skipped": 0, "noise_filtered": 0,
        "dedup_skipped": 0, "entities_deferred": 0, "budget_interrupted": False,
    }
    seen_file_ids: set[str] = set()
    run_start = datetime.now(timezone.utc)

    for i, item in enumerate(work):
        folder_id   = item["folder_id"]
        folder_name = item["folder_name"]
        entity      = item["entity"]
        watermark_key = item["watermark_key"]

        # Between-entity budget boundary: stop cleanly, defer the rest.
        if deadline is not None and time.monotonic() >= deadline:
            aggregate["budget_interrupted"] = True
            aggregate["entities_deferred"] = len(work) - i
            log.info(
                "founders_os: budget reached before %s — deferring %d entity/ies "
                "to next run", entity, aggregate["entities_deferred"],
            )
            break

        log.info("founders_os: sweeping %s -> entity=%s", folder_name, entity)

        cutoff = datetime.now(timezone.utc) - timedelta(days=freshness_days)
        if item["watermark_ts"] is not None:
            wm_dt = datetime.fromtimestamp(item["watermark_ts"], tz=timezone.utc)
            cutoff = max(cutoff, wm_dt)
        cutoff_str = cutoff.strftime("%Y-%m-%dT%H:%M:%S")

        # D-051 review fix: on COMPLETION the watermark must advance to the time
        # this multi-run sweep BEGAN, not the completing run's start. A resumed
        # run skips the file-processing of already-`completed` folders (perf), so
        # a file dropped into such a folder AFTER it was checkpointed but BEFORE
        # the sweep finished is never re-enumerated on the completing run; if the
        # watermark then jumped to the (later) completing-run start, that file
        # would fall permanently below every future incremental cutoff and be
        # silently lost. Pinning the watermark to the ORIGINAL sweep start keeps
        # such files above the next cutoff (re-enumerated, idempotent upsert).
        # The marker is written EAGERLY (before any tree work) so it survives even
        # a hard SIGKILL that outran the self-budget. Single-run entities pay one
        # extra checkpoint write+delete — negligible.
        start_key = f"founders_os_startmark_{watermark_key}"
        effective_start = int(run_start.timestamp())
        if not dry_run:
            try:
                sm = kb.get_checkpoint(start_key)
                if sm and isinstance(sm.get("started_at"), int):
                    effective_start = sm["started_at"]   # resuming — keep original
                else:
                    kb.set_checkpoint(start_key, {"started_at": effective_start})
            except Exception as exc:  # noqa: BLE001 — never block a sweep on this
                log.warning("founders_os: start-mark failed for %s: %s", entity, exc)

        is_lex         = entity == "LEX"
        score_threshold = _LEX_SCORE_THRESHOLD if is_lex else _DEFAULT_SCORE_THRESHOLD
        stats: dict = {
            "files_enumerated": 0, "files_extracted": 0, "chunks_ingested": 0,
            "phi_skipped": 0, "noise_filtered": 0, "dedup_skipped": 0,
        }

        # Track every checkpoint key this entity touches, and whether the whole
        # entity completed. The watermark advances (and checkpoints clear) ONLY
        # when entity_completed stays True across every subtree.
        entity_completed = True
        ckpt_keys: list[str] = []

        sub_entity_map = _ENTITY_SUB_MAP.get(entity)
        if sub_entity_map:
            subfolders = _list_subfolders(service, folder_id)
            matched = [
                (sf, sub_entity_map[sf["name"].lower()])
                for sf in subfolders if sf["name"].lower() in sub_entity_map
            ]
            # The root tree must NOT re-walk the sub-entity folders (they're each
            # swept as their own sub-entity). This was harmless before only because
            # in-run seen_file_ids dedup skipped them; under resume (fresh
            # seen_file_ids) the root tree would otherwise re-tag sub-entity files
            # with the PARENT entity — a firewall regression. Skip them explicitly.
            sub_skip = frozenset(sf["id"] for sf, _ in matched)
            for sf, sub_entity in matched:
                log.info("founders_os: sweeping sub-entity %s", sub_entity)
                key = f"founders_os_ckpt_{sf['id']}"
                ckpt_keys.append(key)
                done = _sweep_folder_tree(
                    service=service, folder_id=sf["id"],
                    entity=entity, sub_entity=sub_entity,
                    kb=kb, anthropic_client=anthropic_client,
                    cutoff_str=cutoff_str, dry_run=dry_run,
                    is_lex=is_lex, score_threshold=score_threshold,
                    seen_file_ids=seen_file_ids, stats=stats,
                    checkpoint_key=key,
                    sheets_service=sheets_service,
                    deadline_monotonic=deadline,
                )
                if not done:
                    entity_completed = False
                    break
            if entity_completed:
                key = f"founders_os_ckpt_{folder_id}_root"
                ckpt_keys.append(key)
                entity_completed = _sweep_folder_tree(
                    service=service, folder_id=folder_id,
                    entity=entity, sub_entity=None,
                    kb=kb, anthropic_client=anthropic_client,
                    cutoff_str=cutoff_str, dry_run=dry_run,
                    is_lex=is_lex, score_threshold=score_threshold,
                    seen_file_ids=seen_file_ids, stats=stats,
                    checkpoint_key=key,
                    sheets_service=sheets_service,
                    deadline_monotonic=deadline,
                    skip_folder_ids=sub_skip,
                )
        else:
            key = f"founders_os_ckpt_{folder_id}"
            ckpt_keys.append(key)
            entity_completed = _sweep_folder_tree(
                service=service, folder_id=folder_id,
                entity=entity, sub_entity=None,
                kb=kb, anthropic_client=anthropic_client,
                cutoff_str=cutoff_str, dry_run=dry_run,
                is_lex=is_lex, score_threshold=score_threshold,
                seen_file_ids=seen_file_ids, stats=stats,
                checkpoint_key=key,
                sheets_service=sheets_service,
                deadline_monotonic=deadline,
            )

        aggregate["entities_swept"] += 1
        for k in ("files_enumerated", "files_extracted", "chunks_ingested",
                  "phi_skipped", "noise_filtered", "dedup_skipped"):
            aggregate[k] += stats.get(k, 0)

        if entity_completed:
            if not dry_run:
                try:
                    # effective_start = the ORIGINAL sweep start (see start-mark
                    # note above), so mid-sweep file drops are not skipped past.
                    kb.set_sync_state(watermark_key, effective_start)
                except Exception as exc:
                    log.warning("founders_os: watermark advance failed for %s: %s", entity, exc)
                # Entity fully swept — clear its resume checkpoints + start marker.
                for key in (*ckpt_keys, start_key):
                    try:
                        kb.delete_checkpoint(key)
                    except Exception as exc:  # noqa: BLE001
                        log.warning("founders_os: checkpoint clear failed for %s: %s", key, exc)
            log.info(
                "founders_os: %s done -- enumerated=%d extracted=%d ingested=%d "
                "phi=%d noise=%d dedup=%d",
                entity, stats["files_enumerated"], stats["files_extracted"],
                stats["chunks_ingested"], stats["phi_skipped"],
                stats["noise_filtered"], stats["dedup_skipped"],
            )
        else:
            # Budget cut mid-entity: watermark stays put, checkpoints persisted
            # for resume next run. Stop here — the budget is spent.
            aggregate["budget_interrupted"] = True
            aggregate["entities_deferred"] = len(work) - i - 1
            log.info(
                "founders_os: %s INTERRUPTED by budget (checkpointed for resume) "
                "-- ingested=%d so far; %d later entity/ies deferred",
                entity, stats["chunks_ingested"], aggregate["entities_deferred"],
            )
            break

    log.info(
        "founders_os: COMPLETE -- entities=%d enumerated=%d extracted=%d "
        "ingested=%d phi=%d noise=%d dedup=%d deferred=%d%s",
        aggregate["entities_swept"], aggregate["files_enumerated"],
        aggregate["files_extracted"], aggregate["chunks_ingested"],
        aggregate["phi_skipped"], aggregate["noise_filtered"],
        aggregate["dedup_skipped"], aggregate["entities_deferred"],
        " [budget-interrupted]" if aggregate["budget_interrupted"] else "",
    )
    return aggregate
